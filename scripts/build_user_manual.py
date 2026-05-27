from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "展位销售管理系统用户使用手册.docx"

FONT = "Microsoft YaHei"
FONT_EAST_ASIA = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 39, 50)
MUTED = RGBColor(92, 103, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "AAB7C4"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_style(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_style(footer, after=0)
    run = footer.add_run("展位销售管理系统用户使用手册")
    set_run_font(run, 8, color=MUTED)


def title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_style(p, before=110, after=12, line=1.15)
    r = p.add_run("展位销售管理系统")
    set_run_font(r, 26, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_style(p, after=26)
    r = p.add_run("用户使用手册")
    set_run_font(r, 18, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_style(p, after=18)
    r = p.add_run("适用于超级管理员、管理员、业务员和企业账号")
    set_run_font(r, 11, color=MUTED)

    add_info_table(
        doc,
        [
            ("系统定位", "展会展位销售、客户保护、订单收款、审核流和企业展务资料填报的一体化管理系统。"),
            ("主要对象", "会务/销售管理人员、业务员、参展企业联系人。"),
            ("默认入口", "登录页先选择展会类别与登录展会，再输入账号和密码。"),
            ("生成日期", datetime.now().strftime("%Y-%m-%d")),
        ],
    )

    add_callout(
        doc,
        "使用前建议",
        "首次部署后，请先由超级管理员检查展会信息、销售规则、展馆展区、账号部门和工作日同步，再开始导入或绘制展位、创建业务账号。生产环境请及时修改默认管理员密码。",
    )
    doc.add_page_break()


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    set_para_style(p)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        set_para_style(p, after=4)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        r = p.add_run("• ")
        set_run_font(r, 11, color=BLUE)
        r = p.add_run(item)
        set_run_font(r, 11)


def add_steps(doc, items):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        set_para_style(p, after=4)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        r = p.add_run(f"{idx}. ")
        set_run_font(r, 11, bold=True, color=BLUE)
        r = p.add_run(item)
        set_run_font(r, 11)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    set_table_borders(table, color="D7E1EA")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    set_para_style(p, after=3)
    r = p.add_run(title)
    set_run_font(r, 11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph(text)
    set_para_style(p, after=0)
    for run in p.runs:
        set_run_font(run, 10)
    doc.add_paragraph()


def add_info_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2700, 6660])
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "说明"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            set_para_style(p, after=0)
            for run in p.runs:
                set_run_font(run, 10, bold=True, color=DARK_BLUE)
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_para_style(p, after=0)
                for run in p.runs:
                    set_run_font(run, 10, bold=(idx == 0), color=DARK_BLUE if idx == 0 else INK)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            set_para_style(p, after=0)
            for run in p.runs:
                set_run_font(run, 9.5, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            for p in cells[idx].paragraphs:
                set_para_style(p, after=0, line=1.18)
                for run in p.runs:
                    set_run_font(run, 9.2)
    doc.add_paragraph()


def build_manual():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)
    title_page(doc)

    add_heading(doc, "1. 手册概览", 1)
    add_body(doc, "本手册按实际业务流程组织，帮助用户从登录、基础配置、展位管理、客户录入、订单收款、审核，到企业展务资料填报和数据导出完成全流程操作。")
    add_matrix(
        doc,
        ["角色", "主要权限", "典型工作"],
        [
            ("超级管理员", "拥有系统最高权限，可管理展会、展会类别、账号、部门、规则、展位图、审核和导出。", "初始化系统、创建展会、分配账号权限、维护基础规则。"),
            ("管理员/管理人员", "可管理当前展会的业务数据、规则、展位图、审核、导出和企业展务。", "配置销售规则、审核水单/合同/展务申请、查看统计。"),
            ("业务员", "可录入客户、认领客户、创建订单、选择展位、上传合同或水单、维护自己的客户与订单。", "客户跟进、展位预留、收款资料提交。"),
            ("企业账号", "仅进入企业后台，填报会刊、参展证、楣板和展具增租信息。", "提交参展资料和展务申请。"),
        ],
        [1700, 3700, 3960],
    )

    add_heading(doc, "2. 登录与首页", 1)
    add_steps(
        doc,
        [
            "打开系统登录页，先选择“展会类别”，再选择“登录展会”。",
            "输入账号和密码，点击“登录”。超级管理员默认账号为 admin，默认密码为 admin123；生产环境应及时修改。",
            "登录后进入数据看板。左侧菜单根据账号角色显示不同功能，右上方可退出登录。",
            "如企业收到免登录链接，可直接通过链接进入企业后台；链接有效期由管理员在销售规则中配置。",
        ],
    )
    add_callout(doc, "登录失败排查", "若提示“请先选择展会类别和展会”，说明未完成展会选择；若提示账号或密码不正确，请检查账号状态、密码和是否使用了正确环境的数据。")

    add_heading(doc, "3. 数据看板", 1)
    add_body(doc, "数据看板用于快速查看当前展会的销售和展位使用情况，管理员可看到全局统计，业务员主要看到自己相关的数据。")
    add_bullets(
        doc,
        [
            "核心指标：展位总量、已预订展位、到账展位、已收款、合计面积、标摊数量、光地面积和未收款金额。",
            "销售概览：展示预订率、到账率、回款率、展位状态分布，以及按部门和业务员拆分的销售明细。",
            "待办提醒：合同审核、水单审核、即将释放、企业资料缺失、楣板/展具审核等事项会集中提醒。",
            "点击部分指标可跳转到参展企业列表并自动筛选相关订单。",
        ],
    )

    add_heading(doc, "4. 客户数据管理", 1)
    add_heading(doc, "4.1 新客户列表", 2)
    add_steps(
        doc,
        [
            "进入“客户数据 > 新客户列表”，点击新增客户。",
            "填写企业名称、简称、联系人、电话、邮箱、地址、税号、境内/境外地区等信息。",
            "系统会根据企业名称、简称、联系人、电话、税号提示相似客户，避免重复录入。",
            "保存后客户进入保护期，并占用业务员或部门的客户保护名额。",
        ],
    )
    add_heading(doc, "4.2 老客户列表与客户公海", 2)
    add_bullets(
        doc,
        [
            "老客户通常来自历史成交或历史参展数据，可用于复购跟进。",
            "客户公海用于存放保护到期、下保或释放的客户；业务员可按权限认领。",
            "管理员可查看全局客户，业务员主要查看自己或部门可访问的客户。",
            "客户详情页会汇总客户资料、订单流程、附件、水单、企业展务进度和操作日志。",
        ],
    )

    add_heading(doc, "5. 展位图与展位管理", 1)
    add_heading(doc, "5.1 展位图管理（管理员）", 2)
    add_steps(
        doc,
        [
            "进入“展位图管理”，可上传展位图底图，并设置图纸宽高、比例尺等参数。",
            "在图上新增或选择展位，维护展位号、展馆、展区、类型、状态、宽深、面积和价格。",
            "支持批量选择展位后进行移动、复制、对齐、贴合、锁定、删除等操作。",
            "可创建外部/内部障碍物，系统会按障碍物影响面积重新计算展位计费面积。",
            "如需快速生成测试或标准网格，可使用生成网格功能；正式使用前应核对展位号和尺寸。",
        ],
    )
    add_heading(doc, "5.2 销售展位图", 2)
    add_bullets(
        doc,
        [
            "销售展位图用于业务员查看展位状态并辅助选位。",
            "展位颜色表示状态：空闲、预留、已成交、停用等。",
            "创建订单或更换展位时，可在弹窗展位图中搜索展位号并点击空闲展位选择。",
            "锁定或停用的展位不能被业务员选择。",
        ],
    )

    add_heading(doc, "6. 订单、收款与成交流程", 1)
    add_steps(
        doc,
        [
            "从新客户或参展入口创建订单，选择订单类型。展位订单需要选择展位；无展位订单需要填写订单标题、金额和明细。",
            "系统根据展位价格、优惠规则、首款比例和预留有效期生成订单金额与到期时间。",
            "业务员按销售流程上传合同或水单。若启用“合同先审”，需合同审核通过后才能上传水单。",
            "管理员在审核页审核合同、水单或销售订单水单。审核通过后计入已收款；驳回时必须填写原因。",
            "达到首款比例或全款条件后，订单进入成交或部分到账状态。成交展位可生成企业账号或免登录链接。",
            "未在期限内完成要求的订单，可由系统释放检查释放展位和客户保护关系。",
        ],
    )
    add_matrix(
        doc,
        ["状态", "含义", "常见下一步"],
        [
            ("空闲", "展位可被选择。", "业务员创建展位订单。"),
            ("已预留", "订单已占用展位，等待合同/水单/收款。", "上传合同或水单，等待审核。"),
            ("待审款", "水单或收款资料已提交，等待管理员审核。", "管理员审核通过或驳回。"),
            ("已首款成交", "收款达到成交判断条件。", "生成企业账号，进入企业展务资料填报。"),
            ("已释放/已取消", "订单或展位占用已结束。", "展位回到可销售池或按业务规则处理。"),
        ],
        [1800, 4100, 3460],
    )

    add_heading(doc, "7. 审核中心", 1)
    add_body(doc, "审核中心由管理员使用，集中处理所有需要确认的材料和变更。审核通过会推动订单或展务流程继续；审核驳回会记录驳回原因并通知相关人员重新处理。")
    add_bullets(
        doc,
        [
            "客户合同审核：审核业务员提交的合同文件。",
            "客户水单审核：审核客户保护流程中的水单文件。",
            "销售订单水单审核：审核订单收款金额和水单附件。",
            "楣板审核：审核企业提交的楣板名称修改。",
            "展具增租审核：审核企业提交的展具增租申请。",
            "订单变更审核：审核更换展位、退订展位、特殊订单等申请。",
            "已审核历史：按流程、结果和关键字查询历史审核记录。",
        ],
    )

    add_heading(doc, "8. 企业展务模块", 1)
    add_heading(doc, "8.1 管理端操作", 2)
    add_bullets(
        doc,
        [
            "管理员或业务员可在成交订单中生成企业账号，或生成免登录链接发送给企业联系人。",
            "企业展务汇总页可查看订单、企业账号、会刊信息、参展证人数、楣板状态和展具增租情况。",
            "对于已生成的企业账号，可重新生成密码；免登录链接有效期受销售规则限制。",
        ],
    )
    add_heading(doc, "8.2 企业端填报", 2)
    add_steps(
        doc,
        [
            "企业通过账号密码或免登录链接进入企业后台。",
            "填写会刊信息，包括企业介绍、产品介绍、宣传视频和产品图片。",
            "提交参展证人员信息，可新增或删除人员记录。",
            "提交楣板修改申请；管理员审核后生效。",
            "选择需要增租的展具和数量，提交后等待管理员审核。",
        ],
    )

    add_heading(doc, "9. 展会、规则与账号配置", 1)
    add_heading(doc, "9.1 展会管理", 2)
    add_bullets(
        doc,
        [
            "展会信息包括展会编号、名称、日期、地点、类别和关联展会。",
            "超级管理员可维护展会类别、新增展会、分配展会权限，并可清理指定展会数据。",
            "至少需要保留一个展会，否则登录页无法完成展会选择。",
        ],
    )
    add_heading(doc, "9.2 销售规则管理", 2)
    add_matrix(
        doc,
        ["配置项", "作用", "建议"],
        [
            ("展位价格", "设置标摊和光地基础单价。", "上线前确认单位和计价口径。"),
            ("首款比例", "判断订单是否达到成交条件。", "常见为 30%，按合同规则设置。"),
            ("预留有效期", "决定订单未提交有效材料时的释放时间。", "可按工作日或自然日计算。"),
            ("销售流程", "支持直接上传水单或合同先审后上传水单。", "合同管控严格时选择合同先审。"),
            ("客户保护", "设置新老客户保护天数和保护名额。", "可按业务员或部门统计。"),
            ("企业链接有效期", "控制免登录链接可访问天数。", "不应超过展会倒计时。"),
            ("驳回原因模板", "统一审核驳回口径。", "维护常用驳回原因，减少自由输入。"),
        ],
        [1900, 4300, 3160],
    )
    add_heading(doc, "9.3 基础资料与账号", 2)
    add_bullets(
        doc,
        [
            "展馆管理：维护展馆名称，供展位图选择。",
            "展区管理：维护展区名称和颜色，销售展位图按展区颜色展示。",
            "展具管理：维护企业可选的展具名称、尺寸、价格、图片和启用状态。",
            "工作日同步：同步中国大陆节假日/调休数据，用于工作日截止时间计算。",
            "国家地区同步：更新境外客户可选国家/地区。",
            "部门管理：新增、修改和删除部门，用于业务员归属和部门目标统计。",
            "账号管理：创建业务员、管理员和超级管理员账号，并给业务员/管理员分配部门。",
        ],
    )

    add_heading(doc, "10. 导出与附件", 1)
    add_bullets(
        doc,
        [
            "导出汇总提供订单 CSV、企业展务 CSV 和附件清单导出。",
            "附件包括合同、水单、图片、视频、展具缩略图等，系统按权限控制访问。",
            "导出的 CSV 可用 Excel 打开；如出现编码提示，建议使用 UTF-8 导入方式打开。",
        ],
    )

    add_heading(doc, "11. 推荐操作顺序", 1)
    add_steps(
        doc,
        [
            "超级管理员登录，确认当前展会和展会类别。",
            "维护展馆、展区、展具、销售规则、工作日同步和国家地区。",
            "创建部门、业务员和管理员账号，配置展会权限。",
            "上传展位图底图，绘制或导入展位，检查展位号、尺寸、价格和状态。",
            "业务员录入客户并创建订单，选择展位后提交合同或水单。",
            "管理员处理审核，订单成交后生成企业账号或免登录链接。",
            "企业填报展务资料，管理员审核楣板和展具增租。",
            "展会执行期间通过数据看板和导出汇总追踪销售、收款和展务进度。",
        ],
    )

    add_heading(doc, "12. 常见问题", 1)
    add_info_table(
        doc,
        [
            ("无法登录", "确认已选择展会类别和展会；检查账号、密码、账号是否启用；超级管理员无需展会授权，但展会必须存在。"),
            ("业务员看不到展会", "检查超级管理员是否为该展会分配了业务员或管理员权限。"),
            ("不能选择展位", "确认展位状态为空闲，且没有被锁定、停用、预留或成交。"),
            ("不能上传水单", "可能存在待审水单、合同先审未通过、订单已全款、或上传期限已过。"),
            ("客户无法新增", "检查必填项、相似客户提示、客户保护名额和是否已有当前展会有效订单。"),
            ("企业链接不可用", "检查订单是否为成交展位订单、链接是否过期、是否被重新生成。"),
            ("导出后乱码", "使用 Excel 的“从文本/CSV 导入”，编码选择 UTF-8。"),
        ],
    )

    add_callout(
        doc,
        "数据安全提醒",
        "格式化或迁移云服务器前，应先备份数据库和 storage/uploads 附件目录。业务数据可清空，但必须保留或重建超级管理员、展会基础数据、工作日同步数据和国家地区配置。",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_manual()
    print(path)
